from bs4 import BeautifulSoup
import requests
import datetime
import time
from datetime import datetime
# from datetime import datetime,date
from time import gmtime, strftime
# import Tkinter
from tkinter import *
import tkinter as tk
from tkinter import ttk
import Pmw, sys
# import tkMessageBox
# import tkFont
from PIL import ImageTk, Image
from multiprocessing import Process
from threading import Thread
import webbrowser
from time import sleep
from tkinter.ttk import Progressbar
from multiprocessing import Pool
import re
from multiprocessing.pool import ThreadPool
import urllib.request
import threading
# import SendNewsByMail
import textwrap

'''                                                        SITES
 #----------------------------------------------------------------------------------------------------------------------------------------------------------------------

 '''

global ref, ref1
ref1 = 0
ref = 0

global cdate, dsearch
dsearch = 0

now = datetime.now()
cdate = now.strftime("%m %d %Y")
month = (cdate.split()[0])
day = (cdate.split()[1])
year = ((cdate.split()[2]))
cdate = year + month + day
cdate = int(cdate)


def Date_to_int(dt, site):
    import datetime
    global cdate

    if site == 'dark':
        month = (dt.split('/')[0])
        month = int(month)
        if month < 10:
            month = str(month)
            month = "0" + month
        month = str(month)
        day = (dt.split('/')[1])
        year = ((dt.split('/')[2]))
        dt = str(year) + str(month) + str(day)
        dt = int(dt)

    elif site == 'bbc':
        day = (dt.split()[0])
        month = (dt.split()[1])
        month = (time.strptime(month, "%b").tm_mon)
        if month < 10:
            month = str(month)
            month = "0" + month
        year = '2020'
        dt = str(year) + str(month) + str(day)
        dt = int(dt)

    elif site == 'scoop':
        day = (dt.split()[1])
        day = day[:-1]
        month = (dt.split()[0])
        month = (time.strptime(month, "%b").tm_mon)
        if month < 10:
            month = str(month)
            month = "0" + month
        year = '2020'
        dt = str(year) + str(month) + str(day)
        dt = int(dt)

    elif site == 'krebs':
        day = (dt.split()[1])
        month = (dt.split()[0])
        month = (time.strptime(month, "%b").tm_mon)
        if month < 10:
            month = str(month)
            month = "0" + month
        year = '2020'
        dt = str(year) + str(month) + str(day)
        dt = int(dt)

    elif site == 'vice':
        month = (dt.split('.')[0])
        day = (dt.split('.')[1])
        year = ((dt.split('.')[2]))
        if int(month) < 10:
            month = str(month)
            month = "0" + month
        dt = "20" + str(year) + str(month) + str(day)
        dt = int(dt)

    elif site == 'zdnet':
        month = (dt.split('-')[1])
        day = (dt.split('-')[2])
        year = ((dt.split('-')[0]))
        dt = year + month + day
        dt = int(dt)

    elif site == 'sweek':
        month = (dt[0])
        month = int((datetime.datetime.strptime(month, '%B').month))
        if month < 10:
            month = str(month)
            month = "0" + month
        day = (dt[1])
        day = (day[:-1])
        year = ((dt[2]))
        dt = str(year) + str(month) + str(day)
        dt = int(dt)

    elif site == 'wired':
        month = (dt.split('.')[0])
        day = (dt.split('.')[1])
        year = ((dt.split('.')[2]))
        dt = "20" + str(year) + str(month) + str(day)
        dt = int(dt)

    else:
        month = (dt.split()[0])
        month = int((datetime.datetime.strptime(month, '%B').month))
        if month < 10:
            month = str(month)
            month = "0" + month
        day = (dt.split()[1])
        day = (day[:-1])
        year = ((dt.split()[2]))
        dt = str(year) + str(month) + str(day)
        dt = int(dt)

    return (dt)


def bleep(url):
    res = requests.get(url)  # make request to url
    bleep_soup = BeautifulSoup(res.text, 'html.parser')  # bs4 digest
    bleep_latest_news = str(bleep_soup.find_all(class_='tab-pane'))  # save latest news text as variable
    bleep_soup2 = BeautifulSoup(bleep_latest_news, 'html.parser')  # gather text from saved variable

    bleep_titles = []  # empty lists to store titles
    bleep_links = []
    bleep_date = []  # empty list to store links
    dt = []

    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')

    for items in bleep_soup.find_all('h4'):
        for item in items.find_all('a'):
            bleep_links.append(item.get("href"))
            bleep_titles.append(item.string)

    for items in bleep_soup.find_all('li'):
        for items2 in items.find_all('div', {'class': 'bc_latest_news_text'}):
            for items1 in items.find_all('ul'):
                for item in items1.find_all('li', {'class': 'bc_news_date'}):
                    bleep_date.append(item.string)
                    site = 'threat'
                    dt.append(Date_to_int(item.string, site))

    x = "News.txt"
    file_to_insert_text = open(x, "a")

    i = 0
    for d in bleep_date:
        if (cdate - dsearch) <= dt[i]:
            file_to_insert_text.write("\n")
            file_to_insert_text.write(bleep_titles[i])
            file_to_insert_text.write(", ")
            file_to_insert_text.write(d)
            file_to_insert_text.write("\n")
            file_to_insert_text.write(bleep_links[i])
            file_to_insert_text.write("\n")
            i += 1
        else:
            break

    return x


# 2.2. cyberscoop security news function
def cyberscoop(url):
    cyberscoop_page = requests.get(url)  # page request
    cyberscoop_soup = BeautifulSoup(cyberscoop_page.text, 'html.parser')  # bs4 digest

    if url != "https://www.fedscoop.com/attend/#upcoming-virtual-events":
        x = "News.txt"
        file_to_insert_text = open(x, "a")
        dates = []
        dt = []
        titles = []
        links = []
        for items in cyberscoop_soup.find_all('strong'):
            for item in items.find_all('a'):
                date = (item.next_sibling.string)
                date = date[3:]
                dates.append(date)
                site = 'scoop'
                dt.append(Date_to_int(date, site))

        for items in cyberscoop_soup.find_all('a', class_='article-thumb__title'):  # search loop
            titles.append(items.get_text().strip())  # print title
            links.append(items.get('href'))  # print link

        i = 0
        for d in dates:
            if (cdate - dsearch) <= dt[i]:
                file_to_insert_text.write("\n")
                file_to_insert_text.write(titles[i])
                file_to_insert_text.write(", ")
                file_to_insert_text.write(d)
                file_to_insert_text.write("\n")
                file_to_insert_text.write(links[i])
                file_to_insert_text.write("\n")
            i += 1
    else:
        headlines = []
        links = []
        date = []
        x = "Events.txt"
        file_to_insert_text = open(x, "a")
        res = requests.get(url)
        soup = BeautifulSoup(res.text, 'lxml')
        i = 0
        for item2 in soup.find_all('div', {'class': 'events__list events__list--home'}):
            for items in item2.find_all('div', {'class': 'event-thumb__information'}):
                for item in items.find_all('a', {"class": "event-thumb__title"}):
                    links.append(item.get("href"))
                    for item1 in item.find_all("h3"):
                        headlines.append(item1.string)
            break
        for items in soup.find_all("div", {"class": "date"}):
            for item in items.find_all("span"):
                date.append(item.string)

        i = 0
        for h in headlines:
            file_to_insert_text.write("\n")
            file_to_insert_text.write(h[:-20])
            file_to_insert_text.write(",   ")
            file_to_insert_text.write(date[i])
            # file_to_insert_text.write("\n\n")
            # file_to_insert_text.write(inner_body[i + 3].text)
            file_to_insert_text.write("\n")
            file_to_insert_text.write(links[i])
            # file_to_insert_text.write("\n\n")
            i += 1

    return x


# 2.3. krebs on security news function
def krebs(url):
    krebs_page = requests.get(url)  # make request to url
    krebs_soup = BeautifulSoup(krebs_page.text, 'html.parser')  # bs4 digest
    links = []
    titles = []

    for items in krebs_soup.find_all('a', rel='bookmark'):  # search loop
        titles.append(items.get_text())  # article title
        links.append(items.get('href'))  # article link

    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')
    dt = []
    dates = []

    for item1 in soup.find_all(("small")):
        for items in item1.find_all('br'):
            date1 = (items.next_sibling.string)
            dates.append(date1)
            site = 'krebs'
            dt.append(Date_to_int(date1, site))

    x = "News.txt"
    file_to_insert_text = open(x, "a")

    i = 0
    for d in dates:
        if (cdate - dsearch) <= dt[i]:
            file_to_insert_text.write("\n")
            file_to_insert_text.write(titles[i])
            file_to_insert_text.write(", ")
            file_to_insert_text.write(d)
            file_to_insert_text.write("\n")
            file_to_insert_text.write(links[i])
            file_to_insert_text.write("\n")
            i += 1
        else:
            break


# 2.4. motherboard security news function
def vice(url):
    mother_page = requests.get(url)  # page request
    mother_soup = BeautifulSoup(mother_page.text, 'html.parser')  # bs4 digest
    titles = []
    links = []
    for items in mother_soup.find_all('a', class_='topics-card__heading-link'):
        titles.append(items.get_text().strip())  # print title
        links.append(items.get('href'))  # print link

    dates = []
    dt = []
    dif = []
    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')
    for items in soup.find_all('div', {'class': "ff--accent size--6 uppercased"}):
        date = items.string
        dates.append(date)
        site = 'vice'
        if date[-1] != "o":
            dt.append(Date_to_int(date, site))
        else:
            if (((date.split(' ')[1])) == "day") or (((date.split(' ')[1])) == "days"):
                dif.append(int((date.split(' ')[0])))
            else:
                dif.append(0)
    x = "News.txt"
    file_to_insert_text = open(x, "a")
    i = 0
    j = 0
    for d in dates:
        if (i >= len(dif)):
            if (cdate - dsearch) <= dt[j]:
                file_to_insert_text.write("\n")
                file_to_insert_text.write(titles[i])
                file_to_insert_text.write(", ")
                file_to_insert_text.write(d)
                file_to_insert_text.write("\n")
                file_to_insert_text.write(links[i])
                file_to_insert_text.write("\n")
            else:
                break
            j += 1
        else:
            if (dsearch) >= dif[i]:
                file_to_insert_text.write("\n")
                file_to_insert_text.write(titles[i])
                file_to_insert_text.write(", ")
                file_to_insert_text.write(d)
                file_to_insert_text.write("\n")
                file_to_insert_text.write(links[i])
                file_to_insert_text.write("\n")
            else:
                break
        i += 1

    return x


# 2.5. Kapersky Securelist security news function
def securelist(url):
    seclist_page = requests.get(url)  # page request
    seclist_soup = BeautifulSoup(seclist_page.text, 'html.parser')  # bs4 digest
    titles = []
    links = []
    for items in seclist_soup.find_all(class_='entry-title'):  # search loop
        a_tag = items.a  # scrape a tags
        a_title = a_tag.get('title')  # scrape title
        a_link = a_tag.get('href')  # scrape links
        titles.append(a_title)
        links.append(a_link)
    dates = []
    dt = []
    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')
    for items in soup.find_all('div', {'class': "entry-time"}):
        for item in items.find_all('time'):
            date = item.string[:-10]
            dates.append(date)
            site = 'scoop'
            dt.append(Date_to_int(date, site))

    x = "News.txt"
    file_to_insert_text = open(x, "a")
    i = 0
    for d in dates:
        if (cdate - dsearch) <= dt[i]:
            file_to_insert_text.write("\n")
            file_to_insert_text.write(titles[i])
            file_to_insert_text.write(", ")
            file_to_insert_text.write(d)
            file_to_insert_text.write("\n")
            file_to_insert_text.write(links[i])
            file_to_insert_text.write("\n")
        else:
            break
        i += 1
    return x


# 2.6. zdnet security news function
def zdnet(url):
    zdnet_page = requests.get(url)  # page request
    zdnet_soup = BeautifulSoup(zdnet_page.text, 'html.parser')  # bs4 digest
    zdnet_latest_news = str(zdnet_soup.find_all('div', id='latest'))  # latest news text
    bleep_soup2 = BeautifulSoup(zdnet_latest_news, 'html.parser')  # bs4 digest of latest news

    titles = []
    links = []
    for items in bleep_soup2.find_all('a', class_='thumb'):  # search loop
        titles.append(items.get('title'))  # print titles
        links.append("https://www.zdnet.com" + items.get('href'))

    dates = []
    dt = []
    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')
    for items in soup.find_all('p', {'class': "meta"}):
        for item in items.find_all("span"):
            date = item.get("data-date")
            date = date[:-9]
            dates.append(date)
            site = 'zdnet'
            dt.append(Date_to_int(date, site))

    x = "News.txt"
    file_to_insert_text = open(x, "a")
    i = 0
    for d in dates:
        if (cdate - dsearch) <= dt[i] and (i < 15):
            file_to_insert_text.write("\n")
            file_to_insert_text.write(titles[i])
            file_to_insert_text.write(", ")
            file_to_insert_text.write(d)
            file_to_insert_text.write("\n")
            file_to_insert_text.write(links[i])
            file_to_insert_text.write("\n")
        else:
            break
        i += 1
    return x


def CywareNews(url):
    site = 'cyware'
    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')
    headlines = soup.findAll('h1', {'class': 'cy-card__title m-0 cursor-pointer pb-3'})  # all headline object
    span = soup.findAll('div', {'class': 'd-flex align-items-center pb-2'})
    descr = soup.findAll('div', {'class': 'cy-card__description'})
    links = soup.findAll('a')
    if url == "https://cyware.com/cyber-security-news-articles":
        x = 'Cyware.txt'
    elif url == 'https://cyware.com/alerts/filter?alert_type=A&category_slug=mobile-security-news':
        x = 'Mobile.txt'
    elif url == 'https://cyware.com/alerts/filter?alert_type=A&category_slug=social-media-threats':
        x = 'SocialMedia.txt'
    elif url == 'https://cyware.com/category/malware-and-vulnerabilities-news':
        x = 'Malware.txt'
    elif url == "https://cyware.com/alerts/filter?alert_type=A&category_slug=interesting-tweets":
        x = "Tweets.txt"
    cnt = 0
    i = 1
    d = []
    for x1 in descr:
        d.append(x1.string)
        print(x1.string)
    links = []
    for news in links:
        link = news.get("href")
        print(link)
        if link is not None:
            if (link[:5] == "https"):
                links.append(link)
                print(link)
            i += 1
    headlines = []
    for news in headlines:
        h = news.string
        print(h)
        if (x == "Tweets.txt"):
            headlines.append(h[18:])
        else:
            headlines.append(h)
    dates = []
    for news in span:
        dt = news.find('span', {'class': 'cy-card__meta'})
        date = dt.string
        print(date)
        dates.append(date)
        cnt += 1
    file_to_insert_text = open(x, "a")
    file_to_insert_text.write("\n")
    for i in range(cnt):
        file_to_insert_text.write(dates[i])
        file_to_insert_text.write("\n")
    return x


def COVIDNews(url):
    site = 'covid'
    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')
    if url == "https://cyware.com/blog/live-updates-covid-19-cybersecurity-alerts-b313":
        x = 'COVID-19.txt'
    file_to_insert_text = open(x, "a")  # file that I will write into

    i = 0
    j = 1
    cnt = 0
    news_article = soup.findAll('span')
    dates = []
    new = []
    file_to_insert_text.write("\n")
    for news in news_article:
        headline = news.string
        if (headline is not None):
            if ((j % 3 == 0) and (j > 3)):
                dates.append(headline)
                cnt += 1
            if ((i > 3) and (i % 3 == 0)):
                news.append(headline)
                headline = textwrap.fill(headline, 120)
                file_to_insert_text.write(headline)
                file_to_insert_text.write("\n")
                file_to_insert_text.write(dates[cnt - 1])
                file_to_insert_text.write("\n\n")
            i += 1
            j += 1
        if (i > 50):
            break
    return x


global temph
temph = []
for i in range(10):
    temph.append(" ")


def ThreatPostNews(url):
    global temph, ref

    site = 'threat'
    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')
    date = soup.findAll('time', )
    if url == "https://threatpost.com":
        x = 'News.txt'
    elif url == 'https://threatpost.com/category/cloud-security/':
        x = 'Cloud.txt'
    elif url == 'https://threatpost.com/category/malware-2/':
        x = 'Malware.txt'
    elif url == "https://threatpost.com/category/vulnerabilities/":
        x = 'Vulnerabilities.txt'
    elif url == "https://threatpost.com/category/mobile-security/":
        x = 'Mobile.txt'
    file_to_insert_text = open(x, "a")
    i = 1
    c = 1
    for news in soup.findAll('h2', {'class': 'c-card__title'}):
        links = news.find('a', )
        links = links.get('href')
        headline = news.string
        if (c == 1):
            if (temph[0] == headline) and (ref == 1):
                break
            else:
                temph[0] = headline
        dt = date[i].string
        dt = Date_to_int(dt, site)
        if (i <= 4 or i > 9):
            if (cdate - dsearch) <= dt:
                file_to_insert_text.write("\n")
                file_to_insert_text.write(headline)
                file_to_insert_text.write(", ")
                file_to_insert_text.write(date[i].string)
                file_to_insert_text.write("\n")
                file_to_insert_text.write(links)
                file_to_insert_text.write("\n")
            i += 1
        if dt < (cdate - dsearch) and (i <= 5 or i > 10):
            break
        c += 1
    return x


def BBCCyberNews(url):
    global temph, ref
    site = 'bbc'
    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')
    news_box = soup.find('div', {'class': 'gel-layout__item gel-3/5@l'})
    news_article = news_box.find_all('article')  # The top headline objects
    date = soup.findAll('span', {'class': 'qa-meta-date gs-u-mr gs-u-display-inline-block'})
    x = 'News.txt'
    file_to_insert_text = open(x, "a")  # file that I will write into
    # file_to_insert_text.write(("                 Latest Cyber News: \n\n"))
    links = soup.findAll('a', {'class': 'qa-heading-link lx-stream-post__header-link'})
    i = 1
    c = 1
    for news in news_article:
        headline = news.find('header')
        if (c == 1):
            if (temph[1] == headline) and (ref == 1):
                break
            else:
                temph[1] = headline
        dt = date[i - 1].string
        dt = Date_to_int(dt, site)
        if (cdate - dsearch) <= dt:
            file_to_insert_text.write("\n")
            file_to_insert_text.write(headline.text)
            file_to_insert_text.write(", ")
            date1 = (date[i - 1].string) + ' 2020'
            file_to_insert_text.write(date1)
            body = news.find('div', {'class': 'gel-layout__item gel-5/8@l'})
            file_to_insert_text.write("\n")
            file_to_insert_text.write("https://www.bbc.com" + links[i - 1].get('href'))
            file_to_insert_text.write("\n")
        i += 1
        if dt < (cdate - dsearch):
            break
        c += 1
    return x


def TheHackerNews(url):
    global temph, ref
    site = 'hack'
    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')
    date = soup.findAll('div', {'class': 'item-label'})
    d1 = soup.findAll('i', {'class': 'icon-font icon-calendar'})
    date = []
    for d2 in d1:
        date.append(d2.next_sibling)

    if url == "https://thehackernews.com/":
        x = 'News.txt'
    elif url == 'https://thehackernews.com/search/label/Malware':
        x = 'Malware.txt'
    elif url == "https://thehackernews.com/search/label/Vulnerability":
        x = 'Vulnerabilities.txt'
    # elif url == "https://threatpost.com/category/mobile-security/":
    #     x = 'Mobile.txt'
    # elif url == 'https://threatpost.com/category/cloud-security/':
    #     x = 'Cloud.txt'

    file_to_insert_text = open(x, "a")
    links = soup.findAll('a', {'class': 'story-link'})
    i = 1
    c = 1
    for news in soup.findAll('h2', {'class': 'home-title'}):
        dt = date[i - 1].string
        dt = Date_to_int(dt, site)
        headline = news.string
        if (c == 1):
            if (temph[2] == headline) and (ref == 1):
                break
            else:
                temph[2] = headline
        if (cdate - dsearch) <= dt:
            file_to_insert_text.write("\n")
            file_to_insert_text.write(headline)
            file_to_insert_text.write(", ")
            file_to_insert_text.write(date[i - 1].string)
            file_to_insert_text.write("\n")
            file_to_insert_text.write(links[i - 1].get('href'))
            file_to_insert_text.write("\n")
        i += 1
        if dt < (cdate - dsearch):
            break
        c += 1
    return x


def DarkReadNews(url):
    global temph, ref
    site = 'dark'
    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')
    d1 = soup.findAll('span', {'class': 'allcaps smaller'})
    date = []
    j = 0
    for d2 in d1:
        date.append(d2.next_sibling)
        date[j] = date[j].split()
        date[j] = date[j][-1]
        j += 1
    if url == "https://www.darkreading.com/":
        x = 'News.txt'
    elif url == 'https://www.darkreading.com/cloud-security.asp':
        x = 'Cloud.txt'
    elif url == "https://www.darkreading.com/vulnerabilities-threats.asp":
        x = 'Vulnerabilities.txt'
    file_to_insert_text = open(x, "a")
    i = 1
    c = 1
    for n1 in soup.findAll('header', {'class': 'strong medium'}):
        news = n1.find('a', )
        headline = news.get('title')
        if (c == 1):
            if (temph[3] == headline) and (ref == 1):
                break
            else:
                temph[3] = headline
        links = news.get('href')
        dt = date[i - 1]
        dt = Date_to_int(dt, site)
        if (cdate - dsearch) <= dt:
            file_to_insert_text.write("\n")
            file_to_insert_text.write(headline)
            file_to_insert_text.write(", ")
            file_to_insert_text.write(date[i - 1])
            file_to_insert_text.write("\n")
            file_to_insert_text.write("https://www.darkreading.com/" + links)
            file_to_insert_text.write("\n")
        i += 1
        if dt < (cdate - dsearch):
            break
        c += 1
    return x


def SecWeekNews(url):
    global temph, ref

    site = 'sweek'
    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')

    if url == "https://www.securityweek.com/cybercrime":
        x = 'News.txt'
    elif url == 'https://www.securityweek.com/virus-threats':
        x = 'Malware.txt'
    elif url == "https://www.securityweek.com/mobile-wireless":
        x = 'Mobile.txt'

    file_to_insert_text = open(x, "a")
    i = 1
    c = 1
    temp = soup.findAll('div', {'class': 'views-field-title'})
    for n1 in temp:
        news = n1.find('span', {'class': 'field-content'})
        new = news.find("a", )
        headline = new.string
        if (c == 1):
            if (temph[4] == headline) and (ref == 1):
                break
            else:
                temph[4] = headline
        links = news.find('a', )
        links = links.get('href')
        res = requests.get("https://www.securityweek.com/" + links)
        soup = BeautifulSoup(res.text, 'lxml')
        temp = soup.findAll('div', {'class': 'submitted'})
        for n1 in temp:
            date = n1.find('div')
            date = date.text
            date = date.split()[-3:]
            dt = Date_to_int(date, site)
            date1 = date[0] + " " + date[1] + " " + date[2]
        if (cdate - dsearch) <= dt:
            file_to_insert_text.write("\n")
            file_to_insert_text.write(headline)
            file_to_insert_text.write(", ")
            file_to_insert_text.write(date1)
            file_to_insert_text.write("\n")
            file_to_insert_text.write("https://www.securityweek.com/" + links)
            file_to_insert_text.write("\n")
        i += 1
        if dt < (cdate - dsearch):
            break
        c += 1
    return x


def SecMagazine(url):
    global temph, ref

    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')
    date = soup.findAll('div', {'class': 'date article-summary__post-date'})

    if url == "https://www.securitymagazine.com/topics/2236-cyber-security-news":
        x = 'News.txt'
    # elif url == 'https://thehackernews.com/search/label/Malware':
    #    x = 'Malware.txt'
    # elif url == "https://thehackernews.com/search/label/Vulnerability":
    #    x = 'Vulnerabilities.txt'
    #

    file_to_insert_text = open(x, "a")
    i = 1
    c = 1
    for news in soup.findAll('h2', {'class': 'headline article-summary__headline'}):
        new = news.find("a", )
        links = news.find('a', )
        links = links.get('href')
        dt = date[i - 1].string
        dt = Date_to_int(dt, 'smag')
        headline = new.string
        if (c == 1):
            if (temph[5] == headline) and (ref == 1):
                break
            else:
                temph[5] = headline
        if (cdate - dsearch) <= dt:
            file_to_insert_text.write("\n")
            file_to_insert_text.write(headline)
            file_to_insert_text.write(", ")
            file_to_insert_text.write(date[i - 1].string)
            file_to_insert_text.write("\n")
            file_to_insert_text.write(links)
            file_to_insert_text.write("\n")
        i += 1
        if (cdate - dsearch) > dt:
            break
        c += 1
    return x


def Wired(url):
    global temph, ref

    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'lxml')

    if url == "https://www.wired.com/tag/cybersecurity/":
        x = 'News.txt'

    file_to_insert_text = open(x, "a")
    i = 1
    c = 1
    for news in soup.find_all('h2'):
        headline = news.string
        if (c == 1):
            if (temph[6] == headline) and (ref == 1):
                break
            else:
                temph[6] = headline
        links = news.parent
        links = links.get("href")

        res = requests.get("https://www.wired.com/" + links)
        soup = BeautifulSoup(res.text, 'lxml')
        temp = soup.find('time')
        date = temp.string
        dt = Date_to_int(date, 'wired')
        if (cdate - dsearch) <= dt:
            file_to_insert_text.write("\n")
            file_to_insert_text.write(headline)
            file_to_insert_text.write(", ")
            file_to_insert_text.write(date)
            file_to_insert_text.write("\n")
            file_to_insert_text.write("https://www.wired.com/" + links)
            file_to_insert_text.write("\n")
        i += 1
        if (cdate - dsearch) > dt:
            break
        c += 1
    return x


'''                                                        CRAWLERS
 #----------------------------------------------------------------------------------------------------------------------------------------------------------------------

 '''


def Crawl():
    global temp, c
    c = 0
    temp = 0

    def Crawl1():
        global check, temp, ref, crawl1_count
        check = 'News.txt'
        # NEWS
        f = open('News.txt', 'r+')
        if (ref == 0):
            f.truncate(5)
        crawl1_count = 0

        def Crawl11():
            global crawl1_count
            (BBCCyberNews("https://www.bbc.com/news/topics/cz4pr2gd85qt/cyber-security"))
            crawl1_count += 1

        def Crawl12():
            global crawl1_count
            (ThreatPostNews("https://threatpost.com"))
            crawl1_count += 1

        def Crawl13():
            global crawl1_count
            (TheHackerNews("https://thehackernews.com/"))
            crawl1_count += 1

        def Crawl14():
            global crawl1_count
            (DarkReadNews("https://www.darkreading.com/"))
            crawl1_count += 1

        def Crawl15():
            global crawl1_count
            (SecMagazine("https://www.securitymagazine.com/topics/2236-cyber-security-news"))
            crawl1_count += 1

        def Crawl16():
            global crawl1_count
            (Wired("https://www.wired.com/tag/cybersecurity/"))
            crawl1_count += 1

        def Crawl17():
            global crawl1_count
            (SecWeekNews("https://www.securityweek.com/cybercrime"))
            crawl1_count += 1

        def Crawl18():
            global crawl1_count
            bleep("https://www.bleepingcomputer.com/news/security/")
            crawl1_count += 1

        def Crawl19():
            global crawl1_count
            cyberscoop("https://www.cyberscoop.com/")
            crawl1_count += 1

        def Crawl110():
            global crawl1_count
            krebs("https://krebsonsecurity.com")
            crawl1_count += 1

        def Crawl111():
            global crawl1_count
            vice("https://motherboard.vice.com/en_us/topic/hacking")
            crawl1_count += 1

        def Crawl112():
            global crawl1_count
            securelist("https://securelist.com/")
            crawl1_count += 1

        def Crawl113():
            global crawl1_count
            zdnet("https://www.zdnet.com/topic/security")
            crawl1_count += 1

        Thread(target=Crawl11).start()
        Thread(target=Crawl12).start()
        Thread(target=Crawl13).start()
        Thread(target=Crawl14).start()
        Thread(target=Crawl15).start()
        Thread(target=Crawl16).start()
        Thread(target=Crawl17).start()
        Thread(target=Crawl18).start()
        Thread(target=Crawl19).start()
        Thread(target=Crawl110).start()
        Thread(target=Crawl111).start()
        Thread(target=Crawl112).start()
        Thread(target=Crawl113).start()

        while crawl1_count < 12:
            pass

        # (ScMagazine("https://www.scmagazine.com/home/security-news/"))

        temp += 1

    def Crawl2():
        # COVID-19
        global temp
        f = open('COVID-19.txt', 'r+')
        f.truncate(9)
        (COVIDNews("https://cyware.com/blog/live-updates-covid-19-cybersecurity-alerts-b313"))
        temp += 1

    def Crawl3():
        # MALWARE
        global temp
        f = open('Malware.txt', 'r+')
        f.truncate(8)
        (ThreatPostNews("https://threatpost.com/category/malware-2/"))
        (TheHackerNews("https://thehackernews.com/search/label/Malware"))
        (SecWeekNews("https://www.securityweek.com/virus-threats"))
        temp += 1

    def Crawl4():
        # VULNERABILITIES
        global temp
        f = open('Vulnerabilities.txt', 'r+')
        f.truncate(16)
        (ThreatPostNews("https://threatpost.com/category/vulnerabilities/"))
        (TheHackerNews("https://thehackernews.com/search/label/Vulnerability"))
        (DarkReadNews("https://www.darkreading.com/vulnerabilities-threats.asp"))
        temp += 1

    def Crawl5():
        # MOBILE
        global temp
        f = open('Mobile.txt', 'r+')
        f.truncate(16)
        (ThreatPostNews("https://threatpost.com/category/mobile-security/"))
        (SecWeekNews("https://www.securityweek.com/mobile-wireless"))
        temp += 1

    def Crawl6():
        global temp
        f = open('Events.txt', 'r+')
        f.truncate(7)
        cyberscoop("https://www.fedscoop.com/attend/#upcoming-virtual-events")
        temp += 1

    def Crawl7():
        # CLOUD
        global temp
        f = open('Cloud.txt', 'r+')
        f.truncate(6)
        (ThreatPostNews("https://threatpost.com/category/cloud-security/"))
        (DarkReadNews("https://www.darkreading.com/cloud-security.asp"))
        temp += 1

    t1 = Thread(target=Crawl1)
    t1.start()
    t2 = Thread(target=Crawl2)
    t2.start()
    t3 = Thread(target=Crawl3)
    t3.start()
    t4 = Thread(target=Crawl4)
    t4.start()
    t5 = Thread(target=Crawl5)
    t5.start()
    t6 = Thread(target=Crawl6)
    t6.start()
    t7 = Thread(target=Crawl7)
    t7.start()

    while temp < 7:
        time.sleep(1)
        pass
        # print(temp, " Threads have finished")
    c = 1


def Combine_files(filenames):
    with open('CyberNews.txt', 'w') as CyberNews:
        for fname in filenames:
            with open(fname) as infile:
                for line in infile:
                    CyberNews.write(line)


'''                                                        BUTTONS
 #----------------------------------------------------------------------------------------------------------------------------------------------------------------------

 '''


def Color(Frame):
    def __init__(self, master=None):
        Frame.__init__(self, master)
        self.master = master
        self.master.configure(background='black')


global ucheck
ucheck = 0


def updated():
    global update, ucheck, u, tempu
    if (ucheck == 1):
        if update == 0:
            pass
        else:
            Refresh()

    if update != 0:
        tu = threading.Timer(u, updated).start()
        ucheck = 1


def Refresh():
    global win, des, ref, ref1, dsearch, update
    if (ref1 == dsearch):
        ref = 1
    else:
        ref = 0
    ref1 = dsearch

    des = 0
    # find a way to check if there are new alerts
    win = tk.Toplevel(root)

    def disable_event():
        pass

    win.protocol("WM_DELETE_WINDOW", disable_event)

    window_height = 200
    window_width = 200
    screen_width = win.winfo_screenwidth()
    screen_height = win.winfo_screenheight()
    x_cordinate = int((screen_width / 2) - (window_width / 2))
    y_cordinate = int((screen_height / 2) - (window_height / 2))
    win.geometry("{}x{}+{}+{}".format(window_width, window_height, x_cordinate, y_cordinate))
    win.resizable(False, False)
    win.configure(background='black')

    def c2():
        global win, des
        Crawl()
        Click()
        des = 1

    def c1():
        global des

        def load2(l):
            global des
            loading = tk.Label(win, text=l, width=40, pady=10)
            loading.place(relx=0, rely=0, relwidth=1, relheight=1)

            time.sleep(0.3)

        def load1():
            global p, c, des
            while des == 0:
                load2('REFRESHING.')
                if des == 1:
                    break
                load2('REFRESHING..')
                if des == 1:
                    break
                load2('REFRESHING...')
                if des == 1:
                    break
                load2('REFRESHING.')
                if des == 1:
                    break
                load2('REFRESHING..')
                if des == 1:
                    break
                load2('REFRESHING...')
                if des == 1:
                    break
            win.destroy()

        load1()

    cc1 = Thread(target=c1)
    cc1.start()
    cc2 = Thread(target=c2)
    cc2.start()

    win.mainloop()


def About():
    global check
    check = 'About.txt'
    Click()


def Save():
    global check
    from functools import partial

    def Save2():
        global savefile
        tkWindow.withdraw()
        savefile = s2.get()

        from docx import Document
        doc = Document()
        file = check
        with open(file, 'r') as openfile:
            line = openfile.read()
            doc.add_paragraph(line)
            doc.save("TEXTS\\" + savefile + ".docx")

        def pdf():
            from docx2pdf import convert
            convert("TEXTS\\" + savefile + ".docx", "C:\\Users\\User\\Desktop\\" + savefile + ".pdf")

        pdf()

        return

    tkWindow = Tk()
    tkWindow.resizable(False, False)
    window_height = 150
    window_width = 400
    screen_width = tkWindow.winfo_screenwidth()
    screen_height = tkWindow.winfo_screenheight()
    x_cordinate = int((screen_width / 2) - (window_width / 2))
    y_cordinate = int((screen_height / 2) - (window_height / 2))
    tkWindow.geometry("{}x{}+{}+{}".format(window_width, window_height, x_cordinate, y_cordinate))
    tkWindow.title('Save')

    s1 = Label(tkWindow, text="Save as")
    s1.place(relx=0.2, rely=0.4, relwidth=0.2, relheight=0.2)
    s1.focus_set()

    s2 = Entry(tkWindow, )
    s2.place(relx=0.37, rely=0.4, relwidth=0.3, relheight=0.2)
    s2.focus_set()

    Save = partial(Save2)
    SubmitButton = Button(tkWindow, text="Save", command=Save)
    SubmitButton.place(relx=0.4, rely=0.7, relwidth=0.2, relheight=0.3)
    SubmitButton.focus_set()

    tkWindow.mainloop()


def SendMail():
    import datetime
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders

    def Send(fromaddr, password, toaddr):
        global check

        tkWindow = Tk()
        window_height = 150
        window_width = 400
        screen_width = tkWindow.winfo_screenwidth()
        screen_height = tkWindow.winfo_screenheight()
        x_cordinate = int((screen_width / 2) - (window_width / 2))
        y_cordinate = int((screen_height / 2) - (window_height / 2))
        tkWindow.geometry("{}x{}+{}+{}".format(window_width, window_height, x_cordinate, y_cordinate))
        tkWindow.title('Send Email')

        l1 = tk.Label(root, text="Cyber News Insider", fg="white", bg="black", pady=12)
        Success = Label(tkWindow, width=50, text="Succesfully sent")
        Success.config(font=('calibri', 20))
        Fail = Label(tkWindow, width=50, text="Fail to sent")
        Fail.config(font=('calibri', 20))

        try:

            # instance of MIMEMultipart
            msg = MIMEMultipart()
            msg['From'] = fromaddr
            msg['To'] = toaddr

            now = datetime.datetime.now()
            Subject = "Cyber News For: " + (now.strftime('%m/%d/%Y'))

            msg['Subject'] = Subject
            body = "Daily Cyber News"
            msg.attach(MIMEText(body, 'plain'))

            filename = check
            from docx import Document
            import glob
            import os
            import docx

            doc = Document()
            file = check
            i = 1

            with open(file, 'r') as openfile:
                line = openfile.read()
                doc.add_paragraph(line)
                doc.save(
                    "TEXTS\\" + file[:-4] + ".docx")
            tcheck = file[:-4] + ".docx"

            def pdf():
                from docx2pdf import convert

                convert("TEXTS\\" + tcheck)

            pdf()

            tcheck2 = file[:-4] + ".pdf"
            attachment = open(
                "TEXTS\\" + tcheck2, "rb")
            p = MIMEBase('application', 'octet-stream')
            p.set_payload((attachment).read())

            encoders.encode_base64(p)
            p.add_header('Content-Disposition', "attachment; filename= %s" % tcheck2)
            msg.attach(p)

            # creates SMTP session
            s = smtplib.SMTP('smtp.gmail.com', 587)
            s.starttls()
            s.login(fromaddr, password)

            text = msg.as_string()  # Converts the Multipart msg into a string

            s.sendmail(fromaddr, toaddr, text)  # send the mail
            Success.pack(side='top', fill=tk.X)
            s.quit()

        except:
            Fail.pack(side='top', fill=tk.X)

        tkWindow.mainloop()

    pass

    def Credentials():

        from functools import partial

        def Submit2():
            tkWindow.withdraw()
            mail = mailEntry.get()
            password = passwordEntry.get()
            mail2 = mail2Entry.get()

            Send(mail, password, mail2)

            return

        tkWindow = Tk()
        window_height = 150
        window_width = 400
        screen_width = tkWindow.winfo_screenwidth()
        screen_height = tkWindow.winfo_screenheight()
        x_cordinate = int((screen_width / 2) - (window_width / 2))
        y_cordinate = int((screen_height / 2) - (window_height / 2))
        tkWindow.geometry("{}x{}+{}+{}".format(window_width, window_height, x_cordinate, y_cordinate))
        tkWindow.title('Send Email')

        # mail label and text entry box
        mailLabel = Label(tkWindow, text="Your Email").grid(row=0, column=0)
        mailEntry = Entry(tkWindow, textvariable="mail")
        mailEntry.grid(row=0, column=1)

        # password label and password entry box
        passwordLabel = Label(tkWindow, text="Your Password").grid(row=2, column=0)
        passwordEntry = Entry(tkWindow, textvariable="password", show='*')
        passwordEntry.grid(row=2, column=1)

        # mail2 label and text entry box
        mail2dLabel = Label(tkWindow, text="Send to").grid(row=4, column=0)
        mail2Entry = Entry(tkWindow, textvariable="mail2")
        mail2Entry.grid(row=4, column=1)

        Submit = partial(Submit2)
        SubmitButton = Button(tkWindow, text="Submit", command=Submit).grid(row=6, column=1)

        tkWindow.mainloop()

    Credentials()


def Latest():
    global check
    check = "Cyware.txt"
    Click()


def Click():
    global check, show

    # LINK AS TEXT
    def callback1(event):
        webbrowser.open_new(event.widget.cget("text"))

    def callback2(url):
        webbrowser.Chrome.open_new(url)

    f = open(check, "r").read()
    with open(check) as f:
        content = f.readlines()
    # you may also want to remove whitespace characters like `\n` at the end of each line
    content = [x.strip() for x in content]

    canvas = tk.Canvas(root)

    scrolly = tk.Scrollbar(root, orient='vertical', command=canvas.yview)
    # display labels in the canvas
    title = tk.Label(canvas, text=content[0])
    title.configure(font=("Calibri", 12, "bold"))
    canvas.create_window(0, 0, anchor='nw', window=title, height=20)
    for i in range(1, (len(content))):
        if (i == 3 or (i % 3 == 0 and i > 3)) and (check != 'COVID-19.txt' and (check != 'About.txt')):
            label = tk.Label(canvas, text=content[i], cursor='hand2')
            label.bind("<Button-1>", callback1)
        else:
            label = tk.Label(canvas, text=content[i])
            label.configure(font=("Calibri", 11, "bold"))
        canvas.create_window(0, i * 25, anchor='nw', window=label, height=20)
    canvas.configure(scrollregion=canvas.bbox('all'), yscrollcommand=scrolly.set)

    def on_mousewheel(event):
        shift = (event.state & 0x1) != 0
        scroll = -1 if event.delta > 0 else 1
        if shift:
            canvas.xview_scroll(scroll, "units")
        else:
            canvas.yview_scroll(scroll, "units")

    canvas.bind_all("<MouseWheel>", on_mousewheel)

    canvas.place(relx=0, rely=0.129, relwidth=0.89, relheight=0.8)
    scrolly.place(relx=0.87, rely=0.13, relwidth=0.02, relheight=0.799)


def News():
    global check
    check = 'News.txt'
    if show == 1:
        Click()


def Covid():
    global check
    check = 'COVID-19.txt'
    if show == 1:
        Click()


def Malw():
    global check
    check = 'Malware.txt'
    if show == 1:
        Click()


def Vuln():
    global check
    check = 'Vulnerabilities.txt'
    if show == 1:
        Click()


def Mobile():
    global check
    check = 'Mobile.txt'
    if show == 1:
        Click()


def Social():
    global check
    check = 'SocialMedia.txt'
    if show == 1:
        Click()


def Events():
    global check
    check = 'Events.txt'
    if show == 1:
        Click()


def Cloud():
    global check
    check = 'Cloud.txt'
    if show == 1:
        Click()


def updates():
    global update, dseacrh, red1, ref, u, ucheck, tempu
    update = 0

    if b6["text"] == "10min":
        b6["text"] = "Off"
        update = 0

    else:
        b6["text"] = "Off"
        b6["text"] = "10min"
        update = 10

    u = update * 60.0
    ucheck = 0
    updated()


def Days():
    global dsearch, ref, ref1
    dsearch = 0
    if b8["text"] == "Today":
        b8["text"] = "Yesterday"
        dsearch = 1

    elif b8["text"] == "Yesterday":
        b8["text"] = "2 Days ago"
        dsearch = 2

    elif b8["text"] == "2 Days ago":
        b8["text"] = "3 Days ago"
        dsearch = 3

    elif b8["text"] == "3 Days ago":
        b8["text"] = "4 Days ago"
        dsearch = 4

    elif b8["text"] == "4 Days ago":
        b8["text"] = "5 Days ago"
        dsearch = 5

    elif b8["text"] == "5 Days ago":
        b8["text"] = "6 Days ago"
        dsearch = 6

    elif b8["text"] == "6 Days ago":
        b8["text"] = " Weekly "
        dsearch = 7

    else:
        b8["text"] = "Today"
        dsearch = 0


def bar1(x, l):
    global p

    p['value'] = x
    p.pack(pady=170)

    frame.update()
    time.sleep(0.5)


def bar():
    global p, c, d
    while c == 0:
        bar1(0, 'LOADING.')
        bar1(20, 'LOADING..')
        bar1(40, 'LOADING...')
        bar1(60, 'LOADING.')
        bar1(80, 'LOADING..')
        bar1(100, 'LOADING...')
    d = 1
    frame.quit()


'''                                                        MAIN
 #----------------------------------------------------------------------------------------------------------------------------------------------------------------------

 '''


def load():
    global frame, p, d, c
    d = 0
    c = 0
    frame = tk.Toplevel()
    frame.title("LOADING, PLEASE WAIT")
    frame.resizable(False, False)
    window_height = 400
    window_width = 400
    screen_width = frame.winfo_screenwidth()
    screen_height = frame.winfo_screenheight()
    x_cordinate = int((screen_width / 2) - (window_width / 2))
    y_cordinate = int((screen_height / 2) - (window_height / 2))
    frame.geometry("{}x{}+{}+{}".format(window_width, window_height, x_cordinate, y_cordinate))

    def close_program():
        frame.destroy()

    def disable_event():
        pass

    frame.protocol("WM_DELETE_WINDOW", disable_event)

    image = Image.open('header.png')
    image1 = ImageTk.PhotoImage(image)
    loading = tk.Label(frame, image=image1)
    loading.place(relx=0, rely=0, relwidth=1, relheight=1)

    p = ttk.Progressbar(frame, orient=HORIZONTAL, length=200, mode="determinate", takefocus=True, maximum=100)
    p.pack()

    t1 = Thread(target=Crawl)
    t1.start()
    t2 = Thread(target=bar)
    t2.start()

    frame.mainloop()


if __name__ == "__main__":

    root = tk.Tk()
    root.withdraw()
    load()
    frame.withdraw()

    root.title("CNI")
    root.resizable(False, False)  # This code helps to disable windows from resizing
    window_height = 600
    window_width = 1000
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x_cordinate = int((screen_width / 2) - (window_width / 2))
    y_cordinate = int((screen_height / 2) - (window_height / 2))
    root.geometry("{}x{}+{}+{}".format(window_width, window_height, x_cordinate, y_cordinate))

    image = Image.open('Cybermain1.png')
    photo_image = ImageTk.PhotoImage(image)
    label = tk.Label(root, image=photo_image)
    label.place(relx=0.0, rely=0.0, relwidth=1, relheight=1)

    # BUTTONS AND LABELS
    show = 1

    b1 = tk.Button(root, text="Get Latest Cyber News", command=News, pady=8, padx=50)
    b1.pack(side='bottom')

    b5 = tk.Button(root, text='Refresh', command=Refresh, width=20, pady=10).place(relx=0.9, rely=0.24, relwidth=0.1,
                                                                                   relheight=0.1)
    b8 = tk.Button(root, text='Today', command=Days, width=20, pady=10)
    b8.place(relx=0.9, rely=0.38, relwidth=0.1, relheight=0.1)
    b4 = tk.Button(root, text='Save', command=Save, width=20, pady=10).place(relx=0.9, rely=0.52, relwidth=0.1,
                                                                             relheight=0.1)
    b3 = tk.Button(root, text='Send Email', command=SendMail, width=20, pady=10).place(relx=0.9, rely=0.66,
                                                                                       relwidth=0.1, relheight=0.1)
    b7 = tk.Button(root, text='About', command=About, width=20, pady=10).place(relx=0.9, rely=0.8, relwidth=0.1,
                                                                               relheight=0.1)
    l2 = tk.Label(root, text='Get updates : ', width=20, pady=10).place(relx=0.745, rely=0.94, relwidth=0.15,
                                                                        relheight=0.06)
    b6 = tk.Button(root, text="Off", command=updates, width=20, pady=10)
    b6.place(relx=0.9, rely=0.94, relwidth=0.1, relheight=0.06)

    # TYPE OF NEWS
    tk.Button(root, text='News', pady=10, command=News).place(relx=0, rely=0.08, relwidth=0.088, relheight=0.05)
    tk.Button(root, text='Malware', pady=10, command=Malw).place(relx=0.09, rely=0.08, relwidth=0.088, relheight=0.05)
    tk.Button(root, text='Vulnerabilities', pady=10, command=Vuln).place(relx=0.18, rely=0.08, relwidth=0.088,
                                                                         relheight=0.05)
    tk.Button(root, text='Mobile Security', pady=10, command=Mobile).place(relx=0.27, rely=0.08, relwidth=0.088,
                                                                           relheight=0.05)
    tk.Button(root, text='Cloud Security', pady=10, command=Cloud).place(relx=0.36, rely=0.08, relwidth=0.088,
                                                                         relheight=0.05)
    tk.Button(root, text='Events', pady=10, command=Events).place(relx=0.45, rely=0.08, relwidth=0.088, relheight=0.05)
    tk.Button(root, text='COVID-19', pady=10, command=Covid).place(relx=0.54, rely=0.08, relwidth=0.088, relheight=0.05)

    # Search
    edit = Entry(root, )
    edit.place(relx=0.9, rely=0.12, relwidth=0.1, relheight=0.05)
    edit.focus_set()
    butt = tk.Button(root, text='Search')


    def find():
        global check
        f = open('Results.txt', "w")
        f.write('RESULTS')
        f.write("\n")
        f.close()
        f = open('Results.txt', "a")
        text = "News.txt"
        s = edit.get()

        res = []
        i = 0
        temp = ""
        with open(text) as search:
            for line in search:
                if (i % 3) != 0:
                    temp = line.rstrip()
                else:
                    line = line.rstrip()  # remove '\n' at end of line
                    if s.lower() in temp.lower() and s != "":
                        f.write("\n")
                        f.write(temp)
                        f.write("\n")
                        f.write(line)
                        f.write("\n")
                i = i + 1

        f.close()
        check = 'Results.txt'
        Click()


    butt.place(relx=0.9, rely=0.08, relwidth=0.1, relheight=0.05)
    text = Text(root)
    butt.config(command=find)

    root.deiconify()

    root.mainloop()







